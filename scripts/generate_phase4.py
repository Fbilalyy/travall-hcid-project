#!/usr/bin/env python3
"""
IE48L Phase 4: Information Architecture and Interaction Design
TravAll - Travel Super App
Generates a .docx deliverable with:
  1. Task Analysis (3 tasks, one per persona)
  2. System Map (visual diagram)
  3. Screen List (detailed table)
  4. Screen Designs / Wireframes
"""

import os
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOC_PATH = os.path.join(OUT_DIR, "IE48L-Project Phase 4.docx")

# ── Colour palette ──────────────────────────────────────────────
NAVY   = RGBColor(0x00, 0x2B, 0x5C)
TEAL   = RGBColor(0x00, 0x7B, 0x7F)
BLUE   = RGBColor(0x1A, 0x73, 0xE8)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GRAY   = RGBColor(0x5D, 0x6D, 0x7E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper functions ────────────────────────────────────────────
def set_cell_shading(cell, hex_color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="002B5C"):
    """Create a formatted table with colored header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F4F7")

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

def add_heading_styled(doc, text, level=1):
    """Add a heading with custom styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def add_body_text(doc, text, bold=False, italic=False, size=11):
    """Add a paragraph with styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = GRAY
    return p

def fig_to_docx_image(doc, fig, width=6.5):
    """Insert a matplotlib figure into the document."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=200, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return last_paragraph


# ═══════════════════════════════════════════════════════════════
# 1.  SYSTEM MAP DIAGRAM
# ═══════════════════════════════════════════════════════════════
def create_system_map():
    """Create a visual system/site map for TravAll with clear separated columns."""
    fig, ax = plt.subplots(1, 1, figsize=(18, 14))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 14)
    ax.axis('off')
    ax.set_facecolor('white')

    c_root = '#002B5C'
    c_tab  = '#1A73E8'
    c_sub  = '#007B7F'
    c_leaf = '#E67E22'
    c_modal = '#95A5A6'

    def draw_box(x, y, w, h, text, color, fontsize=7.5, text_color='white'):
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                             boxstyle="round,pad=0.08", facecolor=color,
                             edgecolor='#34495E', linewidth=1.0)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
                fontweight='bold', color=text_color, linespacing=1.1)
        return (x, y)

    def arrow(sx, sy, ex, ey, color='#7F8C8D'):
        ax.annotate('', xy=(ex, ey), xytext=(sx, sy),
                    arrowprops=dict(arrowstyle='->', color=color, lw=1.1))

    # Row heights  (top to bottom)
    Y_ROOT = 13.0
    Y_TAB  = 11.2
    Y_SUB  = 9.2
    Y_LEAF = 7.2
    Y_DEEP = 5.2

    # Column centers for 6 sections (Onboarding + 5 tabs)
    # Give each section its own horizontal band so nothing overlaps
    CX = {
        'onb': 1.5,
        'dash': 4.5,
        'visa': 7.8,
        'trip': 11.2,
        'comm': 14.5,
        'prof': 17.0,
    }

    # ── ROOT ──
    root = draw_box(9, Y_ROOT, 3.2, 0.65, 'TravAll App', c_root, 11)

    # ── ONBOARDING (far left) ──
    onb = draw_box(CX['onb'], Y_TAB, 2.0, 0.55, 'Onboarding', c_modal, 8)
    arrow(root[0]-1.2, Y_ROOT-0.33, onb[0]+0.5, Y_TAB+0.27)

    onb1 = draw_box(0.5, Y_SUB, 1.5, 0.45, 'Welcome', c_modal, 7)
    onb2 = draw_box(1.5, Y_LEAF, 1.7, 0.45, 'Passport\nSetup', c_modal, 6.5)
    onb3 = draw_box(2.5, Y_SUB, 1.5, 0.45, 'Preferences', c_modal, 6.5)
    arrow(CX['onb']-0.4, Y_TAB-0.27, 0.5, Y_SUB+0.22)
    arrow(CX['onb'], Y_TAB-0.27, 2.5, Y_SUB+0.22)
    arrow(0.5, Y_SUB-0.22, 1.5, Y_LEAF+0.22)

    # ── 5 MAIN TABS ──
    tabs = [
        ('Dashboard', CX['dash']),
        ('Visa Assistant', CX['visa']),
        ('Trips', CX['trip']),
        ('Community', CX['comm']),
        ('Profile', CX['prof']),
    ]
    tab_nodes = []
    for name, cx in tabs:
        n = draw_box(cx, Y_TAB, 2.2, 0.55, name, c_tab, 8)
        arrow(root[0] + (cx - 9)*0.15, Y_ROOT-0.33, cx, Y_TAB+0.27)
        tab_nodes.append(n)

    # ── DASHBOARD column ──
    dx = CX['dash']
    d1 = draw_box(dx-0.8, Y_SUB, 1.5, 0.45, 'Home Feed', c_sub, 7)
    d2 = draw_box(dx+0.8, Y_SUB, 1.5, 0.45, 'Notifications', c_sub, 6.5)
    arrow(dx-0.3, Y_TAB-0.27, dx-0.8, Y_SUB+0.22)
    arrow(dx+0.3, Y_TAB-0.27, dx+0.8, Y_SUB+0.22)
    d3 = draw_box(dx+0.8, Y_LEAF, 1.5, 0.45, 'Notif. Detail', c_leaf, 7)
    arrow(dx+0.8, Y_SUB-0.22, dx+0.8, Y_LEAF+0.22)

    # ── VISA ASSISTANT column ──
    vx = CX['visa']
    v1 = draw_box(vx-0.8, Y_SUB, 1.4, 0.45, 'Passport\nSelect', c_sub, 6.5)
    v2 = draw_box(vx+0.8, Y_SUB, 1.4, 0.45, 'Destination\nSearch', c_sub, 6.5)
    arrow(vx-0.3, Y_TAB-0.27, vx-0.8, Y_SUB+0.22)
    arrow(vx+0.3, Y_TAB-0.27, vx+0.8, Y_SUB+0.22)

    v3 = draw_box(vx-0.8, Y_LEAF, 1.5, 0.45, 'Visa Detail', c_sub, 7)
    v4 = draw_box(vx+0.8, Y_LEAF, 1.5, 0.45, 'Doc Checklist', c_sub, 6.5)
    arrow(vx-0.8, Y_SUB-0.22, vx-0.8, Y_LEAF+0.22)
    arrow(vx+0.8, Y_SUB-0.22, vx+0.8, Y_LEAF+0.22)

    v5 = draw_box(vx-0.8, Y_DEEP, 1.5, 0.45, 'Doc Upload', c_leaf, 7)
    v6 = draw_box(vx+0.8, Y_DEEP, 1.5, 0.45, 'App Tracker', c_leaf, 7)
    arrow(vx-0.8, Y_LEAF-0.22, vx-0.8, Y_DEEP+0.22)
    arrow(vx+0.8, Y_LEAF-0.22, vx+0.8, Y_DEEP+0.22)

    # ── TRIPS column ──
    tx = CX['trip']
    t1 = draw_box(tx-0.9, Y_SUB, 1.4, 0.45, 'My Trips', c_sub, 7)
    t2 = draw_box(tx+0.9, Y_SUB, 1.4, 0.45, 'Search &\nBook', c_sub, 6.5)
    arrow(tx-0.3, Y_TAB-0.27, tx-0.9, Y_SUB+0.22)
    arrow(tx+0.3, Y_TAB-0.27, tx+0.9, Y_SUB+0.22)

    t3 = draw_box(tx-0.9, Y_LEAF, 1.4, 0.45, 'Trip Detail', c_sub, 7)
    arrow(tx-0.9, Y_SUB-0.22, tx-0.9, Y_LEAF+0.22)

    t4 = draw_box(tx-1.5, Y_DEEP, 1.2, 0.45, 'Flight\nSearch', c_leaf, 6)
    t5 = draw_box(tx-0.2, Y_DEEP, 1.2, 0.45, 'Hotel\nSearch', c_leaf, 6)
    t6 = draw_box(tx+1.1, Y_DEEP, 1.2, 0.45, 'Event\nTickets', c_leaf, 6)
    arrow(tx-0.9, Y_LEAF-0.22, tx-1.5, Y_DEEP+0.22)
    arrow(tx-0.9, Y_LEAF-0.22, tx-0.2, Y_DEEP+0.22)
    arrow(tx+0.9, Y_SUB-0.22, tx+1.1, Y_DEEP+0.22)

    t7 = draw_box(tx-0.9, 3.5, 1.3, 0.45, 'Booking\nConfirm', c_leaf, 6)
    t8 = draw_box(tx+0.5, 3.5, 1.3, 0.45, 'Live\nTracking', c_leaf, 6)
    arrow(tx-1.5, Y_DEEP-0.22, tx-0.9, 3.72)
    arrow(tx-0.2, Y_DEEP-0.22, tx+0.5, 3.72)

    # ── COMMUNITY column ──
    cx = CX['comm']
    c1 = draw_box(cx-0.7, Y_SUB, 1.3, 0.45, 'Feed', c_sub, 7)
    c2 = draw_box(cx+0.7, Y_SUB, 1.3, 0.45, 'Create\nPost', c_sub, 6.5)
    arrow(cx-0.3, Y_TAB-0.27, cx-0.7, Y_SUB+0.22)
    arrow(cx+0.3, Y_TAB-0.27, cx+0.7, Y_SUB+0.22)

    c3 = draw_box(cx-0.7, Y_LEAF, 1.3, 0.45, 'Post Detail', c_leaf, 6.5)
    c4 = draw_box(cx+0.7, Y_LEAF, 1.3, 0.45, 'User Profile', c_leaf, 6.5)
    arrow(cx-0.7, Y_SUB-0.22, cx-0.7, Y_LEAF+0.22)
    arrow(cx+0.7, Y_SUB-0.22, cx+0.7, Y_LEAF+0.22)

    # ── PROFILE column ──
    px = CX['prof']
    p1 = draw_box(px-0.7, Y_SUB, 1.3, 0.45, 'My\nPassports', c_sub, 6.5)
    p2 = draw_box(px+0.7, Y_SUB, 1.3, 0.45, 'Settings', c_sub, 7)
    arrow(px-0.3, Y_TAB-0.27, px-0.7, Y_SUB+0.22)
    arrow(px+0.3, Y_TAB-0.27, px+0.7, Y_SUB+0.22)

    p3 = draw_box(px-0.7, Y_LEAF, 1.3, 0.45, 'Travel\nStats', c_leaf, 6.5)
    p4 = draw_box(px+0.7, Y_LEAF, 1.3, 0.45, 'Digital\nWallet', c_leaf, 6.5)
    arrow(px-0.7, Y_SUB-0.22, px-0.7, Y_LEAF+0.22)
    arrow(px+0.7, Y_SUB-0.22, px+0.7, Y_LEAF+0.22)

    p5 = draw_box(px+0.7, Y_DEEP, 1.3, 0.45, 'Add\nPayment', c_leaf, 6.5)
    arrow(px+0.7, Y_LEAF-0.22, px+0.7, Y_DEEP+0.22)

    # ── LEGEND (bottom-left) ──
    leg_y = 2.0
    legend_items = [
        (c_root, 'Root'),
        (c_tab, 'Main Tab (Bottom Nav)'),
        (c_sub, 'Sub-screen'),
        (c_leaf, 'Detail / Leaf Screen'),
        (c_modal, 'Modal / Onboarding')
    ]
    for i, (color, label) in enumerate(legend_items):
        draw_box(1.2, leg_y - i*0.55, 0.8, 0.35, '', color, 6)
        ax.text(1.9, leg_y - i*0.55, label, ha='left', va='center',
                fontsize=7.5, color='#2C3E50')

    ax.set_title('TravAll System Map', fontsize=15, fontweight='bold',
                 color='#002B5C', pad=8)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════
# 2.  WIREFRAME SCREENS
# ═══════════════════════════════════════════════════════════════
def draw_phone_frame(ax, title=""):
    """Draw a phone frame outline."""
    ax.set_xlim(0, 9)
    ax.set_ylim(0, 16)
    ax.set_aspect('equal')
    ax.axis('off')

    # Phone body
    phone = FancyBboxPatch((0.2, 0.2), 8.6, 15.6,
                           boxstyle="round,pad=0.3", facecolor='#FAFAFA',
                           edgecolor='#2C3E50', linewidth=2)
    ax.add_patch(phone)

    # Status bar
    status = FancyBboxPatch((0.5, 14.8), 8, 0.7,
                            boxstyle="round,pad=0.05", facecolor='#002B5C',
                            edgecolor='none')
    ax.add_patch(status)
    ax.text(4.5, 15.15, '9:41', ha='center', va='center',
            fontsize=7, color='white', fontweight='bold')
    ax.text(7.8, 15.15, '100%', ha='right', va='center',
            fontsize=6, color='white')

    if title:
        ax.text(4.5, 14.3, title, ha='center', va='center',
                fontsize=10, fontweight='bold', color='#002B5C')


def draw_bottom_nav(ax, active=0):
    """Draw bottom navigation bar."""
    nav = FancyBboxPatch((0.5, 0.5), 8, 1.2,
                         boxstyle="round,pad=0.05", facecolor='white',
                         edgecolor='#BDC3C7', linewidth=1)
    ax.add_patch(nav)
    tabs = ['Home', 'Visa', 'Trips', 'Social', 'Profile']
    icons = ['⌂', '✈', '🗓', '👥', '👤']
    for i, (tab, icon) in enumerate(zip(tabs, icons)):
        x = 1.3 + i * 1.6
        color = '#1A73E8' if i == active else '#95A5A6'
        ax.text(x, 1.3, icon, ha='center', va='center', fontsize=9, color=color)
        ax.text(x, 0.85, tab, ha='center', va='center', fontsize=5.5,
                color=color, fontweight='bold' if i == active else 'normal')


def draw_button(ax, x, y, w, h, text, color='#1A73E8', text_color='white', fontsize=7):
    btn = FancyBboxPatch((x, y), w, h,
                         boxstyle="round,pad=0.1", facecolor=color,
                         edgecolor='none')
    ax.add_patch(btn)
    ax.text(x + w/2, y + h/2, text, ha='center', va='center',
            fontsize=fontsize, color=text_color, fontweight='bold')


def draw_card(ax, x, y, w, h, title="", subtitle="", color='white'):
    card = FancyBboxPatch((x, y), w, h,
                          boxstyle="round,pad=0.1", facecolor=color,
                          edgecolor='#E0E0E0', linewidth=1)
    ax.add_patch(card)
    if title:
        ax.text(x + 0.3, y + h - 0.35, title, fontsize=7,
                fontweight='bold', color='#2C3E50')
    if subtitle:
        ax.text(x + 0.3, y + h - 0.7, subtitle, fontsize=5.5, color='#7F8C8D')


def create_wireframe_dashboard():
    fig, ax = plt.subplots(figsize=(4.5, 8))
    draw_phone_frame(ax, "Dashboard")
    draw_bottom_nav(ax, active=0)

    # Greeting
    ax.text(1, 13.5, "Good morning, Yavuz! 👋", fontsize=8,
            fontweight='bold', color='#2C3E50')
    ax.text(1, 13.0, "Your next trip is in 3 days", fontsize=6, color='#7F8C8D')

    # Upcoming trip card
    draw_card(ax, 0.7, 11.0, 7.6, 1.7, "✈  Istanbul → Prague", "Mar 30 · Turkish Airlines TK1767")
    draw_button(ax, 5.5, 11.2, 2.5, 0.6, "View Details", '#007B7F', fontsize=6)

    # Quick Actions
    ax.text(1, 10.3, "Quick Actions", fontsize=7, fontweight='bold', color='#002B5C')
    actions = [("🛂 Visa\nCheck", 0.8), ("🔍 Search\nFlights", 2.8),
               ("🏨 Hotels", 4.8), ("🎫 Events", 6.8)]
    for label, x in actions:
        draw_card(ax, x, 8.8, 1.7, 1.2)
        ax.text(x + 0.85, 9.4, label, ha='center', va='center',
                fontsize=5.5, color='#2C3E50')

    # Notifications
    ax.text(1, 8.2, "Recent Updates", fontsize=7, fontweight='bold', color='#002B5C')
    notifs = [
        ("⚠️  Visa rule change: Czech Republic", "2 hours ago", 7.0),
        ("✅  Flight TK1767 confirmed", "Yesterday", 6.0),
        ("💡  Tip: Pack light for 3-day trips", "2 days ago", 5.0),
    ]
    for text, time, y in notifs:
        draw_card(ax, 0.7, y, 7.6, 0.8)
        ax.text(1, y + 0.5, text, fontsize=5.5, color='#2C3E50')
        ax.text(7.8, y + 0.2, time, fontsize=4.5, color='#95A5A6', ha='right')

    # Smart suggestion
    draw_card(ax, 0.7, 3.5, 7.6, 1.2, "", "", '#E8F5E9')
    ax.text(1, 4.35, "💡 Smart Suggestion", fontsize=6, fontweight='bold', color='#2E7D32')
    ax.text(1, 3.9, "Based on your passport: 47 visa-free destinations", fontsize=5.5, color='#2E7D32')

    # Floating action button
    circle = plt.Circle((7.8, 2.5), 0.4, color='#E67E22', zorder=5)
    ax.add_patch(circle)
    ax.text(7.8, 2.5, "+", ha='center', va='center', fontsize=14,
            color='white', fontweight='bold', zorder=6)

    fig.tight_layout()
    return fig


def create_wireframe_visa():
    fig, ax = plt.subplots(figsize=(4.5, 8))
    draw_phone_frame(ax, "Visa Assistant")
    draw_bottom_nav(ax, active=1)

    # Passport selector
    ax.text(1, 13.2, "Your Passport", fontsize=7, fontweight='bold', color='#002B5C')
    draw_card(ax, 0.7, 12.0, 7.6, 1.0)
    ax.text(1.2, 12.6, "🟢  Green Passport (Hususi)", fontsize=7, color='#2C3E50')
    ax.text(7.5, 12.5, "▼", fontsize=9, color='#7F8C8D', ha='right')

    # Search bar
    search_bg = FancyBboxPatch((0.7, 10.8), 7.6, 0.8,
                                boxstyle="round,pad=0.1", facecolor='#F5F5F5',
                                edgecolor='#E0E0E0')
    ax.add_patch(search_bg)
    ax.text(1.2, 11.2, "🔍  Search destination country...", fontsize=6.5, color='#95A5A6')

    # Visa-free destinations header
    ax.text(1, 10.2, "Visa-Free for Green Passport", fontsize=7,
            fontweight='bold', color='#007B7F')
    ax.text(1, 9.7, "87 countries · No visa required", fontsize=5.5, color='#7F8C8D')

    # Country cards
    countries = [
        ("🇩🇪  Germany", "Visa-free · 90 days", "✅", '#E8F5E9', 8.7),
        ("🇫🇷  France", "Visa-free · 90 days", "✅", '#E8F5E9', 7.6),
        ("🇬🇧  United Kingdom", "eVisa required", "📋", '#FFF3E0', 6.5),
        ("🇺🇸  United States", "Visa required", "⚠️", '#FFEBEE', 5.4),
        ("🇯🇵  Japan", "eVisa · 30 days", "📋", '#FFF3E0', 4.3),
    ]
    for name, status, icon, bg, y in countries:
        draw_card(ax, 0.7, y, 7.6, 0.8, "", "", bg)
        ax.text(1.1, y + 0.5, name, fontsize=6.5, fontweight='bold', color='#2C3E50')
        ax.text(1.1, y + 0.15, status, fontsize=5, color='#7F8C8D')
        ax.text(7.8, y + 0.4, icon, fontsize=9, ha='right')

    # Filter buttons
    filters = ["All", "Visa-Free", "eVisa", "Visa Req."]
    for i, f in enumerate(filters):
        x = 0.7 + i * 2.0
        active = (i == 0)
        c = '#1A73E8' if active else '#F5F5F5'
        tc = 'white' if active else '#2C3E50'
        draw_button(ax, x, 3.3, 1.8, 0.55, f, c, tc, fontsize=5.5)

    fig.tight_layout()
    return fig


def create_wireframe_visa_detail():
    fig, ax = plt.subplots(figsize=(4.5, 8))
    draw_phone_frame(ax, "Czech Republic")
    draw_bottom_nav(ax, active=1)

    # Back arrow
    ax.text(0.8, 14.3, "←", fontsize=12, color='#1A73E8')

    # Country header
    ax.text(4.5, 13.2, "🇨🇿", ha='center', fontsize=20)
    ax.text(4.5, 12.5, "Czech Republic", ha='center', fontsize=9,
            fontweight='bold', color='#2C3E50')

    # Status badge
    draw_button(ax, 2.8, 11.8, 3.4, 0.5, "✅ Visa-Free · 90 days",
                '#2E7D32', fontsize=6)

    # Requirements card
    draw_card(ax, 0.7, 9.5, 7.6, 2.0, "Entry Requirements", "")
    reqs = [
        "✓  Valid passport (6+ months)",
        "✓  Return ticket",
        "✓  Hotel reservation or invitation",
        "✓  Travel insurance (€30K min)",
        "✓  Proof of sufficient funds"
    ]
    for i, r in enumerate(reqs):
        ax.text(1.2, 10.9 - i*0.35, r, fontsize=5.5, color='#2C3E50')

    # Document checklist
    draw_card(ax, 0.7, 7.3, 7.6, 1.8, "Document Checklist", "")
    docs = [
        ("☑  Passport scan", True),
        ("☑  Travel insurance", True),
        ("☐  Hotel booking", False),
        ("☐  Return ticket", False),
    ]
    for i, (d, done) in enumerate(docs):
        color = '#2E7D32' if done else '#E67E22'
        ax.text(1.2, 8.6 - i*0.35, d, fontsize=5.5, color=color)

    # Progress bar
    draw_card(ax, 0.7, 6.3, 7.6, 0.7)
    ax.text(1, 6.75, "Readiness: 50%", fontsize=6, fontweight='bold', color='#002B5C')
    bar_bg = FancyBboxPatch((1, 6.4), 6.5, 0.2, boxstyle="round,pad=0.02",
                             facecolor='#E0E0E0', edgecolor='none')
    ax.add_patch(bar_bg)
    bar_fill = FancyBboxPatch((1, 6.4), 3.25, 0.2, boxstyle="round,pad=0.02",
                               facecolor='#1A73E8', edgecolor='none')
    ax.add_patch(bar_fill)

    # Action buttons
    draw_button(ax, 0.7, 5.2, 3.6, 0.7, "📤 Upload Documents", '#1A73E8', fontsize=6.5)
    draw_button(ax, 4.7, 5.2, 3.6, 0.7, "📋 Full Guide", '#007B7F', fontsize=6.5)

    # Alert
    draw_card(ax, 0.7, 3.8, 7.6, 1.0, "", "", '#FFF3E0')
    ax.text(1.1, 4.45, "⚠️  Rule Update (2 hours ago)", fontsize=6,
            fontweight='bold', color='#E65100')
    ax.text(1.1, 4.05, "Insurance minimum increased to €30K from €15K",
            fontsize=5, color='#E65100')

    fig.tight_layout()
    return fig


def create_wireframe_trips():
    fig, ax = plt.subplots(figsize=(4.5, 8))
    draw_phone_frame(ax, "My Trips")
    draw_bottom_nav(ax, active=2)

    # Tabs
    draw_button(ax, 0.7, 13.0, 2.4, 0.5, "Upcoming", '#1A73E8', fontsize=6.5)
    draw_button(ax, 3.3, 13.0, 2.4, 0.5, "Past", '#F5F5F5', '#2C3E50', fontsize=6.5)
    draw_button(ax, 5.9, 13.0, 2.4, 0.5, "Drafts", '#F5F5F5', '#2C3E50', fontsize=6.5)

    # Trip cards
    trips = [
        ("Prague Adventure", "Mar 30 - Apr 2, 2026", "✈ TK1767 · 🏨 Hotel Rott", 10.7, '#E3F2FD'),
        ("Paris Business", "Apr 15 - Apr 18, 2026", "✈ TK1829 · 🏨 Hilton Arc", 8.7, '#F3E5F5'),
        ("Berlin Weekend", "May 1 - May 3, 2026", "✈ TK1721 · 🏨 Pending", 6.7, '#FFF3E0'),
    ]
    for title, dates, details, y, bg in trips:
        draw_card(ax, 0.7, y, 7.6, 1.7, "", "", bg)
        ax.text(1.1, y + 1.3, title, fontsize=7, fontweight='bold', color='#2C3E50')
        ax.text(1.1, y + 0.9, dates, fontsize=5.5, color='#7F8C8D')
        ax.text(1.1, y + 0.5, details, fontsize=5, color='#546E7A')
        draw_button(ax, 5.8, y + 0.2, 2.2, 0.5, "Details →", '#1A73E8', fontsize=5.5)

    # Bottom: add new trip
    draw_button(ax, 1.5, 5.3, 6, 0.8, "+ Plan New Trip", '#E67E22', fontsize=7)

    # Stats mini
    draw_card(ax, 0.7, 3.5, 7.6, 1.4, "2026 Stats", "")
    stats = ["✈ 3 flights", "🏨 3 hotels", "🌍 3 countries", "📅 12 days"]
    for i, s in enumerate(stats):
        ax.text(1 + i * 2, 3.8, s, fontsize=5.5, color='#546E7A')

    fig.tight_layout()
    return fig


def create_wireframe_community():
    fig, ax = plt.subplots(figsize=(4.5, 8))
    draw_phone_frame(ax, "Community")
    draw_bottom_nav(ax, active=3)

    # Tabs
    draw_button(ax, 0.7, 13.0, 2.4, 0.5, "Feed", '#1A73E8', fontsize=6.5)
    draw_button(ax, 3.3, 13.0, 2.4, 0.5, "Tips", '#F5F5F5', '#2C3E50', fontsize=6.5)
    draw_button(ax, 5.9, 13.0, 2.4, 0.5, "My Posts", '#F5F5F5', '#2C3E50', fontsize=6.5)

    # Post 1
    draw_card(ax, 0.7, 10.5, 7.6, 2.2)
    ax.text(1.1, 12.3, "Nihan K. · Verified Traveler 🔵", fontsize=6,
            fontweight='bold', color='#2C3E50')
    ax.text(1.1, 11.9, "Green passport holders: Czech Republic now requires",
            fontsize=5, color='#2C3E50')
    ax.text(1.1, 11.6, "€30K insurance! Update your docs before traveling.",
            fontsize=5, color='#2C3E50')
    ax.text(1.1, 11.1, "🇨🇿 Prague · 45 min ago", fontsize=4.5, color='#95A5A6')
    ax.text(1.1, 10.7, "❤️ 24   💬 8   🔖 Save", fontsize=5, color='#7F8C8D')

    # Post 2
    draw_card(ax, 0.7, 7.8, 7.6, 2.4)
    ax.text(1.1, 9.8, "Yavuz P. · Student Traveler 🎓", fontsize=6,
            fontweight='bold', color='#2C3E50')
    ax.text(1.1, 9.4, "First time in Prague! The astronomical clock is",
            fontsize=5, color='#2C3E50')
    ax.text(1.1, 9.1, "incredible. Highly recommend the Old Town Square.",
            fontsize=5, color='#2C3E50')
    # Image placeholder
    img = FancyBboxPatch((1.1, 8.0), 6.5, 0.8, boxstyle="round,pad=0.05",
                          facecolor='#E0E0E0', edgecolor='#BDBDBD')
    ax.add_patch(img)
    ax.text(4.35, 8.4, "📷 Photo", ha='center', fontsize=6, color='#95A5A6')

    # Post 3
    draw_card(ax, 0.7, 5.5, 7.6, 2.0)
    ax.text(1.1, 7.1, "Melih Ö. · First-Timer 🌟", fontsize=6,
            fontweight='bold', color='#2C3E50')
    ax.text(1.1, 6.7, "TravAll made my first visa application so easy!",
            fontsize=5, color='#2C3E50')
    ax.text(1.1, 6.4, "Upload docs → checklist → done. No stress at all.",
            fontsize=5, color='#2C3E50')
    ax.text(1.1, 5.9, "❤️ 67   💬 15   🔖 Save", fontsize=5, color='#7F8C8D')

    # FAB
    circle = plt.Circle((7.8, 3.0), 0.4, color='#E67E22', zorder=5)
    ax.add_patch(circle)
    ax.text(7.8, 3.0, "✏", ha='center', va='center', fontsize=10,
            color='white', zorder=6)

    fig.tight_layout()
    return fig


def create_wireframe_profile():
    fig, ax = plt.subplots(figsize=(4.5, 8))
    draw_phone_frame(ax, "Profile")
    draw_bottom_nav(ax, active=4)

    # Avatar
    circle = plt.Circle((4.5, 12.5), 0.7, color='#E0E0E0', zorder=3)
    ax.add_patch(circle)
    ax.text(4.5, 12.5, "YP", ha='center', va='center', fontsize=12,
            fontweight='bold', color='#002B5C', zorder=4)
    ax.text(4.5, 11.5, "Yavuz Polat", ha='center', fontsize=9,
            fontweight='bold', color='#2C3E50')
    ax.text(4.5, 11.1, "Student Traveler · Burgundy Passport",
            ha='center', fontsize=5.5, color='#7F8C8D')

    # Stats row
    stats = [("5", "Trips"), ("3", "Countries"), ("18", "Days Abroad")]
    for i, (num, label) in enumerate(stats):
        x = 1.5 + i * 2.8
        ax.text(x, 10.3, num, ha='center', fontsize=11, fontweight='bold', color='#1A73E8')
        ax.text(x, 9.9, label, ha='center', fontsize=5.5, color='#7F8C8D')

    # Menu items
    menu = [
        ("🛂  My Passports", 8.8),
        ("📊  Travel Statistics", 7.8),
        ("💳  Digital Wallet", 6.8),
        ("🔔  Notification Settings", 5.8),
        ("🌙  Dark Mode", 4.8),
        ("🌐  Language", 3.8),
        ("❓  Help & Support", 2.8),
    ]
    for label, y in menu:
        draw_card(ax, 0.7, y, 7.6, 0.7)
        ax.text(1.2, y + 0.35, label, fontsize=6.5, color='#2C3E50')
        ax.text(7.8, y + 0.35, "›", fontsize=10, color='#BDC3C7', ha='right')

    fig.tight_layout()
    return fig


def create_wireframe_onboarding():
    fig, ax = plt.subplots(figsize=(4.5, 8))
    draw_phone_frame(ax, "")

    # Logo area
    ax.text(4.5, 13.0, "TravAll", ha='center', fontsize=18,
            fontweight='bold', color='#002B5C')
    ax.text(4.5, 12.3, "Your Travel Super App", ha='center',
            fontsize=8, color='#7F8C8D')

    # Illustration placeholder
    ill = FancyBboxPatch((1.5, 8.5), 6, 3.2, boxstyle="round,pad=0.1",
                          facecolor='#E3F2FD', edgecolor='#BBDEFB')
    ax.add_patch(ill)
    ax.text(4.5, 10.1, "✈️ 🌍 🛂", ha='center', fontsize=24)

    # Passport selection
    ax.text(4.5, 7.8, "Select Your Passport Type", ha='center',
            fontsize=8, fontweight='bold', color='#2C3E50')

    passport_types = [
        ("🔴  Ordinary (Bordo)", 6.8, False),
        ("🟢  Special (Green)", 5.9, False),
        ("⚫  Diplomatic (Black)", 5.0, False),
        ("🔵  Service (Gray)", 4.1, False),
    ]
    for label, y, sel in passport_types:
        bg = '#E3F2FD' if sel else 'white'
        draw_card(ax, 1.2, y, 6.6, 0.7, "", "", bg)
        ax.text(1.8, y + 0.35, label, fontsize=6.5, color='#2C3E50')
        circ_color = '#1A73E8' if sel else '#E0E0E0'
        c = plt.Circle((7.2, y + 0.35), 0.15, color=circ_color, zorder=3)
        ax.add_patch(c)

    # Continue button
    draw_button(ax, 1.5, 2.8, 6, 0.8, "Continue →", '#1A73E8', fontsize=8)

    # Skip
    ax.text(4.5, 2.2, "Skip for now", ha='center', fontsize=6, color='#7F8C8D')

    # Progress dots
    for i in range(3):
        c = '#1A73E8' if i == 0 else '#E0E0E0'
        dot = plt.Circle((3.7 + i * 0.5, 1.5), 0.1, color=c)
        ax.add_patch(dot)

    fig.tight_layout()
    return fig


# ═══════════════════════════════════════════════════════════════
# 3.  MAIN DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════
def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ══════════════════════════════════════════════════════════
    # COVER PAGE  (matching Phase 2 / Phase 3 format exactly)
    # ══════════════════════════════════════════════════════════
    for _ in range(2):
        doc.add_paragraph()

    # Bogazici University logo
    logo_path = os.path.join(OUT_DIR, "extracted_img_0.png")
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2.2))

    for _ in range(4):
        doc.add_paragraph()

    # IE48L
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IE48L")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run.bold = True

    doc.add_paragraph()

    # Phase title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phase 4:")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TASK ANALYSIS, SYSTEM MAP")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AND SCREEN DESIGNS")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    run.bold = True

    for _ in range(2):
        doc.add_paragraph()

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("29.03.2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    for _ in range(2):
        doc.add_paragraph()

    # Team members (same format as Phase 2/3)
    members = [
        "Özde Avdan - 2021402126",
        "Nilsu Çallıoğulları - 2021402165",
        "Deniz Taş - 2021402123",
        "Fatih Bilal Yılmaz - 2021402174",
    ]
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, "CONTENTS", 1)
    toc_items = [
        ("1. Introduction", "3"),
        ("2. Task Analysis", "4"),
        ("   2.1 Task 1 – Visa Check & Trip Planning (Yavuz)", "4"),
        ("   2.2 Task 2 – Real-Time Alert & Schedule Optimization (Nihan)", "5"),
        ("   2.3 Task 3 – First-Time Setup & Guided Visa Application (Melih)", "6"),
        ("3. System Map", "8"),
        ("4. Screen List", "9"),
        ("5. Screen Designs (Wireframes)", "13"),
        ("   5.1 Onboarding", "13"),
        ("   5.2 Dashboard", "14"),
        ("   5.3 Visa Assistant", "15"),
        ("   5.4 Visa Detail", "16"),
        ("   5.5 My Trips", "17"),
        ("   5.6 Community", "18"),
        ("   5.7 Profile", "19"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{item}")
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY if not item.startswith("   ") else GRAY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Introduction", 1)

    intro_text = (
        "This document presents Phase 4 of the TravAll project: Task Analysis, "
        "Information Architecture, and Screen Designs. TravAll is an all-inclusive "
        "travel super-app that integrates visa checking (by passport type), "
        "flight/hotel booking, event ticketing, an authenticated community, "
        "travel statistics, and a digital wallet into a single mobile platform."
    )
    add_body_text(doc, intro_text)

    approach_text = (
        "Building upon the personas, scenarios, and storyboards developed in Phase 3, "
        "we perform task analysis for three representative user tasks—one per persona. "
        "We then present the system map showing the overall navigation hierarchy, "
        "a comprehensive screen list detailing every UI element and its function, "
        "and wireframe-level screen designs for the key screens of TravAll."
    )
    add_body_text(doc, approach_text)

    # Design principles from Phase 3
    add_heading_styled(doc, "Design Principles (from Phase 2 & 3 Findings)", 2)

    principles = [
        ("Mobile-First & One-Handed Ergonomics", "62.26% of surveyed users operate their phone with one hand. All primary actions are placed within the lower 60% of the screen (thumb zone)."),
        ("Visual Hierarchy & Summarization", "Icon-driven navigation, color-coded visa status (green = visa-free, orange = eVisa, red = visa required), and card-based layouts reduce cognitive load."),
        ("User Control & Transparency", "94.34% of users want final approval on critical decisions. Every booking and document submission requires explicit user confirmation."),
        ("Minimalist Interface", "73.59% prefer clean, minimal UIs. We limit each screen to one primary action and use progressive disclosure for details."),
        ("Dynamic Notifications", "86.79% want real-time alerts for visa rule changes, flight updates, and community activity. A dedicated notification feed is always accessible from the dashboard."),
    ]
    for title, desc in principles:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = TEAL
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)
        run2.font.color.rgb = GRAY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 2. TASK ANALYSIS
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Task Analysis", 1)

    ta_intro = (
        "We conduct task analysis using the complementary template format recommended "
        "in class (Stanton, 2006). Each task table describes the step-by-step process a "
        "user follows to accomplish a goal, documenting: what the user needs to know, "
        "how they learn it, how the action is performed, and the feedback they receive. "
        "We analyze three tasks, one for each persona from Phase 3."
    )
    add_body_text(doc, ta_intro)

    # ── TASK 1: Yavuz ──
    add_heading_styled(doc, "2.1 Task 1 – Check Visa Requirements & Plan a Trip", 2)
    p = doc.add_paragraph()
    run = p.add_run("Persona: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run2 = p.add_run("Yavuz Polat (24, Male, Burgundy Passport, Efficient Student Traveler)")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    p = doc.add_paragraph()
    run = p.add_run("Context: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run2 = p.add_run(
        "Yavuz is an Erasmus student in Germany planning a weekend trip to Prague. "
        "He wants to quickly check whether he needs a visa (burgundy passport), "
        "find a flight, and add the trip to his plan—all one-handed on the metro."
    )
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    doc.add_paragraph()

    task1_headers = ["Step #", "Steps to be Performed", "What Need to Know",
                     "How Know It", "Way Performed", "Feedback"]
    task1_rows = [
        ["1", "Open TravAll app", "Where the app icon is", "Recognize TravAll icon on home screen", "Tap the TravAll icon on phone home screen", "App opens to Dashboard with greeting and upcoming trip summary"],
        ["2", "Navigate to Visa Assistant", "Where Visa tab is in the navigation", "See labeled icon '✈ Visa' in bottom navigation bar", "Tap 'Visa' tab in bottom nav (thumb-reachable zone)", "Visa Assistant screen loads with passport selector and search bar"],
        ["3", "Confirm passport type", "Which passport is currently selected", "See '🔴 Bordo Passport' pre-selected in dropdown", "Verify the correct passport is shown; tap dropdown if need to change", "Passport type confirmed; destination list updates for burgundy passport rules"],
        ["4", "Search for Czech Republic", "How to search for a specific country", "See search bar with placeholder '🔍 Search destination...'", "Tap search bar, type 'Czech' using on-screen keyboard", "Auto-suggestions appear; 'Czech Republic' shows with '✅ Visa-Free · 90 days' badge"],
        ["5", "View visa requirements detail", "How to see full requirements", "See the country card is tappable (chevron indicator)", "Tap 'Czech Republic' card", "Detail screen opens: entry requirements list, document checklist, readiness progress bar"],
        ["6", "Review requirements and checklist", "What documents are needed", "Read the checklist items with ☑/☐ indicators", "Scroll down to read all requirements and check which docs are already uploaded", "See 'Readiness: 50%' progress bar; passport scan and insurance already ☑"],
        ["7", "Navigate to search flights", "How to switch to trip planning", "See 'Quick Actions' on Dashboard or '+ Plan New Trip' button", "Tap back to Dashboard → tap '🔍 Search Flights' quick action card", "Flight search screen appears with destination pre-filled as Prague"],
        ["8", "Search and select a flight", "Available flights and prices", "See list of flights sorted by price/time with airline logos", "Browse results, tap preferred flight (e.g. TK1767)", "Flight detail expands: departure/arrival times, price, seat selection"],
        ["9", "Confirm booking", "Final price and booking details", "Review summary screen with all details and total price", "Review details → tap 'Confirm Booking' button", "Confirmation screen: '✅ Booking Confirmed!' with booking reference and 'Add to My Trips' option"],
        ["10", "View trip in My Trips", "How to see all planned trips", "Tap 'Trips' tab in bottom navigation", "Tap 'Trips' tab", "My Trips list shows the new Prague trip card with flight, dates, and visa status badge"],
    ]

    add_styled_table(doc, task1_headers, task1_rows,
                     col_widths=[1.0, 2.8, 2.5, 2.5, 3.0, 3.2])

    doc.add_page_break()

    # ── TASK 2: Nihan ──
    add_heading_styled(doc, "2.2 Task 2 – Respond to Real-Time Visa Alert & Optimize Schedule", 2)
    p = doc.add_paragraph()
    run = p.add_run("Persona: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run2 = p.add_run("Nihan Katırcı (38, Female, Green Passport, Business Voyager)")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    p = doc.add_paragraph()
    run = p.add_run("Context: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run2 = p.add_run(
        "Nihan receives a push notification about a visa rule change for her upcoming "
        "Czech Republic business trip. She needs to review the update, check her document "
        "readiness, and use the app's AI voice feature to quickly reschedule a meeting."
    )
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    doc.add_paragraph()

    task2_headers = ["Step #", "Steps to be Performed", "What Need to Know",
                     "How Know It", "Way Performed", "Feedback"]
    task2_rows = [
        ["1", "Receive push notification", "That a visa rule has changed", "Push notification appears: '⚠️ Green Passport Alert: Instant entry rule update for Czech Republic'", "Notification appears on lock screen / notification shade", "Banner with TravAll logo, alert icon, and preview text"],
        ["2", "Open notification", "How to access the alert details", "Tap the notification banner", "Tap notification banner", "App opens directly to Visa Detail page for Czech Republic with alert banner highlighted in orange at top"],
        ["3", "Review the rule change", "What has changed and how it affects her trip", "Read the orange alert card: 'Insurance minimum increased to €30K'", "Scroll to read alert details and updated requirements list", "Alert card shows old vs. new rule; affected documents highlighted in orange"],
        ["4", "Check document readiness", "Whether her current documents still meet requirements", "See document checklist with ☑/☐ status indicators", "Scroll down to Document Checklist section", "Checklist shows '☐ Travel Insurance' now needs updating (marked orange), other docs still ☑"],
        ["5", "Upload updated insurance document", "How to upload a new document", "See '📤 Upload Documents' button", "Tap 'Upload Documents' → select 'Travel Insurance' → choose file from phone gallery or camera", "Upload progress bar → '✅ Document uploaded successfully' confirmation; checklist updates to ☑"],
        ["6", "Verify readiness is 100%", "Overall readiness status", "See progress bar updated", "Check the readiness progress bar", "Progress bar now shows '100%' in green; '✅ You are fully prepared!' message"],
        ["7", "Use AI voice to check schedule", "How to activate voice assistant", "See microphone icon in top-right corner of Dashboard", "Return to Dashboard → tap 🎙 microphone icon → say 'Optimize my flight schedule'", "AI voice responds: 'Schedule optimized in 10 seconds, Nihan. Your TK1829 departure moved to 09:15 to avoid conflict with your 14:00 meeting.'"],
        ["8", "Confirm schedule change", "Whether the change is acceptable", "Voice assistant shows summary card with old vs. new times", "Review the proposed change → tap 'Approve Change' button", "'✅ Schedule updated!' confirmation; calendar and trip detail updated"],
        ["9", "Check travel statistics", "Her time savings and trip efficiency", "Navigate to Profile → Travel Statistics", "Tap 'Profile' tab → tap '📊 Travel Statistics'", "Dashboard shows: 'Time Saved: 20%', trip frequency chart, expense breakdown by category"],
    ]

    add_styled_table(doc, task2_headers, task2_rows,
                     col_widths=[1.0, 2.8, 2.5, 3.0, 3.0, 3.2])

    doc.add_page_break()

    # ── TASK 3: Melih ──
    add_heading_styled(doc, "2.3 Task 3 – First-Time Setup & Guided Visa Application", 2)
    p = doc.add_paragraph()
    run = p.add_run("Persona: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run2 = p.add_run("Melih Özgen (21, Male, Ordinary/Bordo Passport, Anxious First-Timer)")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    p = doc.add_paragraph()
    run = p.add_run("Context: ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run2 = p.add_run(
        "Melih has never traveled abroad. He downloads TravAll, sets up his passport, "
        "discovers visa-friendly destinations, and uses the guided document upload "
        "feature to prepare his first visa application—all with step-by-step reassurance."
    )
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRAY

    doc.add_paragraph()

    task3_headers = ["Step #", "Steps to be Performed", "What Need to Know",
                     "How Know It", "Way Performed", "Feedback"]
    task3_rows = [
        ["1", "Download and open TravAll", "Where to get the app", "Search 'TravAll' in App Store / Google Play", "Download app → tap to open", "Welcome screen with TravAll logo and 'Your Travel Super App' tagline; animated travel icons"],
        ["2", "Start onboarding flow", "How to begin setup", "See prominent 'Get Started' button", "Tap 'Get Started' button", "Onboarding Step 1/3: Passport selection screen with 4 passport type options"],
        ["3", "Select passport type", "Which passport type he has", "See descriptions: 'Ordinary (Bordo)', 'Special (Green)', etc.", "Tap '🔴 Ordinary (Bordo)' option", "Selection highlights in blue; 'Continue →' button activates"],
        ["4", "Set travel preferences", "What preferences to configure", "Step 2/3 screen shows options: interests, preferred airlines, budget range", "Select options (e.g., 'Culture', 'Budget-friendly') → tap 'Continue'", "Progress dots update (2/3); preferences saved"],
        ["5", "Complete onboarding", "Final setup step", "Step 3/3: notification preferences and optional biometric setup", "Toggle on 'Visa Rule Alerts' and 'Flight Updates' → tap 'Start Exploring!'", "Dashboard loads with personalized greeting: 'Welcome, Melih! 👋 Let's plan your first trip.'"],
        ["6", "Explore visa-free destinations", "Where to find easy-to-visit countries", "See '💡 Smart Suggestion' card on Dashboard: '47 visa-free destinations for your passport'", "Tap the Smart Suggestion card", "Visa Assistant opens filtered to 'Visa-Free' tab showing countries with ✅ badge"],
        ["7", "Select a destination (Prague)", "Which countries are suitable", "Browse visa-free list with 90-day stay indicators", "Scroll and tap 'Czech Republic – Visa-free · 90 days'", "Visa Detail screen: requirements list, document checklist, '0% readiness' progress bar"],
        ["8", "Start guided document upload", "How to prepare documents", "See '📤 Upload Documents' button and step-by-step guide prompt", "Tap 'Upload Documents'", "Guided upload wizard opens: Step 1 of 5 – 'Passport Front Page' with camera outline and example overlay"],
        ["9", "Upload passport scan", "How to take a good scan", "See example overlay showing correct passport placement and tips", "Position passport within frame → tap capture button", "'✅ Passport scan uploaded!' → auto-extracted data shown for verification → tap 'Confirm'"],
        ["10", "Upload remaining documents", "Which documents are next", "Wizard advances: Step 2 – Insurance, Step 3 – Bank statement, etc.", "Follow guided prompts for each document, uploading or selecting from gallery", "Each upload shows ✅; progress bar increases: 20% → 40% → 60% → 80% → 100%"],
        ["11", "Review final readiness", "Whether all documents are ready", "See '✅ 100% Ready!' badge and green progress bar", "Review summary screen with all documents checked", "'🎉 You're fully prepared for Prague!' celebration animation; 'Final Approval' button"],
        ["12", "Give final approval", "That this is the last confirmation step", "See 'Final Approval' button (per 94.34% user preference for explicit confirmation)", "Tap 'Final Approval' button", "'✅ Application saved! You can submit when ready.' with 'Share with Community' option"],
    ]

    add_styled_table(doc, task3_headers, task3_rows,
                     col_widths=[1.0, 2.6, 2.3, 3.0, 3.0, 3.2])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 3. SYSTEM MAP
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. System Map", 1)

    sm_text = (
        "The system map below shows the complete navigation hierarchy of TravAll. "
        "The app uses a bottom tab navigation pattern with 5 main sections: "
        "Dashboard, Visa Assistant, Trips, Community, and Profile. An onboarding "
        "flow is presented on first launch. Each section branches into sub-screens "
        "and detail/leaf screens following a maximum depth of 3 levels to maintain "
        "simplicity and reduce navigation overhead."
    )
    add_body_text(doc, sm_text)

    doc.add_paragraph()
    print("  Generating system map diagram...")
    fig_sysmap = create_system_map()
    fig_to_docx_image(doc, fig_sysmap, width=6.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Figure 1: TravAll System Map / Site Map")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 4. SCREEN LIST
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Screen List", 1)

    sl_text = (
        "The table below lists all screens and their primary UI elements. "
        "For each element, we describe its function. Global navigation elements "
        "(bottom tab bar, back button) are present on every primary screen unless noted."
    )
    add_body_text(doc, sl_text)

    doc.add_paragraph()

    screen_list_headers = ["Screen / Element", "Function"]
    screen_list_rows = [
        # Onboarding
        ["Onboarding – Welcome", ""],
        ["  TravAll Logo & Tagline", "Displays brand identity and sets user expectation"],
        ["  'Get Started' Button", "Initiates the 3-step onboarding flow"],
        ["  'Log In' Link", "Navigates to login screen for returning users"],
        ["Onboarding – Passport Setup", ""],
        ["  Passport Type Options (4)", "Radio-select: Ordinary, Special, Diplomatic, Service"],
        ["  'Continue →' Button", "Saves selection and advances to Step 2 (Preferences)"],
        ["  Progress Dots (1/3)", "Shows current position in onboarding flow"],
        ["Onboarding – Preferences", ""],
        ["  Interest Tags", "Multi-select: Culture, Adventure, Business, Beach, Food, etc."],
        ["  Budget Range Slider", "Sets preferred budget range for recommendations"],
        ["  'Continue →' Button", "Advances to Step 3 (Notifications)"],
        ["Onboarding – Notifications", ""],
        ["  Toggle: Visa Rule Alerts", "Enables/disables push notifications for visa changes"],
        ["  Toggle: Flight Updates", "Enables/disables real-time flight status notifications"],
        ["  'Start Exploring!' Button", "Completes onboarding; navigates to Dashboard"],

        # Dashboard
        ["Dashboard (Home)", ""],
        ["  Greeting Banner", "Personalized greeting with user name and next trip countdown"],
        ["  Upcoming Trip Card", "Shows next trip: destination, date, airline, with 'View Details' button"],
        ["  Quick Action Cards (4)", "Tap-to-navigate: Visa Check, Search Flights, Hotels, Events"],
        ["  Recent Updates Feed", "Scrollable list of notifications with timestamp"],
        ["  Smart Suggestion Card", "AI-generated recommendation based on passport and preferences"],
        ["  FAB (+) Button", "Floating action button to quickly create new trip"],
        ["  🎙 Voice Button", "Activates AI voice assistant for hands-free interaction"],
        ["  Bottom Navigation Bar", "5 tabs: Home, Visa, Trips, Social, Profile (persistent)"],

        # Visa Assistant
        ["Visa Assistant", ""],
        ["  Passport Selector Dropdown", "Displays current passport; tap to switch between registered passports"],
        ["  Search Bar", "Text input to search destination countries by name"],
        ["  Filter Chips", "Toggle filters: All, Visa-Free, eVisa, Visa Required"],
        ["  Country Cards (List)", "Each card: flag + name + visa status badge + chevron; tap for detail"],
        ["Visa Detail Screen", ""],
        ["  Country Header + Flag", "Large flag icon, country name, visa status badge"],
        ["  Entry Requirements List", "Checklist: passport validity, return ticket, insurance, funds, etc."],
        ["  Document Checklist", "Interactive checklist with ☑/☐ for each required document"],
        ["  Readiness Progress Bar", "Visual indicator (0–100%) of document completeness"],
        ["  '📤 Upload Documents' Button", "Opens guided upload wizard for document submission"],
        ["  '📋 Full Guide' Button", "Opens detailed step-by-step visa guide"],
        ["  Alert Banner", "Orange card showing recent rule changes (if any)"],

        # Trips
        ["My Trips", ""],
        ["  Tab Bar (Upcoming/Past/Drafts)", "Filters trips by status"],
        ["  Trip Cards (List)", "Each card: trip name, dates, flight/hotel info, colored background"],
        ["  'Details →' Button", "Opens Trip Detail for the selected trip"],
        ["  '+ Plan New Trip' Button", "Initiates new trip creation flow"],
        ["  2026 Stats Mini-Card", "Shows flights, hotels, countries, days for current year"],
        ["Trip Detail", ""],
        ["  Trip Header", "Destination name, dates, visa status badge"],
        ["  Flight Section", "Flight details: airline, times, booking reference; tap to expand"],
        ["  Hotel Section", "Hotel name, address, check-in/out; tap to see on map"],
        ["  Events Section", "List of booked events/museums with time and location"],
        ["  Live Tracking Toggle", "Switch to real-time flight tracking view"],
        ["Flight Search", ""],
        ["  Origin / Destination Fields", "Auto-filled from context; editable with search"],
        ["  Date Picker", "Calendar widget for departure and return dates"],
        ["  Results List", "Cards: airline logo, times, duration, price; sortable"],
        ["Hotel Search", ""],
        ["  Destination / Dates", "Pre-filled from trip; editable"],
        ["  Map View Toggle", "Switch between list and map view of hotels"],
        ["  Hotel Cards", "Photo, name, rating, price per night; tap for detail"],
        ["Booking Confirmation", ""],
        ["  Summary Card", "All booking details with total price"],
        ["  'Confirm Booking' Button", "Explicit confirmation (user control principle)"],
        ["  Success Screen", "✅ animation, booking reference, 'Add to My Trips' button"],

        # Community
        ["Community Feed", ""],
        ["  Tab Bar (Feed/Tips/My Posts)", "Navigates between content types"],
        ["  Post Cards", "User avatar, name, badge, text, optional photo, reactions"],
        ["  Reaction Bar", "❤️ Like, 💬 Comment, 🔖 Save actions"],
        ["  FAB (✏) Button", "Opens Create Post screen"],
        ["Create Post", ""],
        ["  Text Input Area", "Rich text editor for post content"],
        ["  Photo Attachment", "Camera/gallery picker for adding images"],
        ["  Location Tag", "Auto-suggest destination tag (e.g., 🇨🇿 Prague)"],
        ["  'Post' Button", "Publishes post to community feed"],

        # Profile
        ["Profile", ""],
        ["  Avatar + Name + Badge", "User photo, name, traveler type badge"],
        ["  Stats Row", "Trips count, countries visited, total days abroad"],
        ["  Menu Items (7)", "Tappable rows: Passports, Statistics, Wallet, Notifications, Dark Mode, Language, Help"],
        ["My Passports", ""],
        ["  Passport Cards", "Each registered passport with type, number, expiry date"],
        ["  'Add Passport' Button", "Register additional passport"],
        ["Travel Statistics", ""],
        ["  Time Saved %", "AI-calculated efficiency improvement"],
        ["  Trip Frequency Chart", "Bar chart of trips per month"],
        ["  Expense Breakdown", "Pie chart by category (flights, hotels, events)"],
        ["  Country Map", "World map with visited countries highlighted"],
        ["Digital Wallet", ""],
        ["  Balance Display", "Current wallet balance in selected currency"],
        ["  Transaction History", "List of payments with date and amount"],
        ["  'Add Payment Method' Button", "Opens payment method registration"],
        ["Settings", ""],
        ["  Notification Toggles", "Per-category notification controls"],
        ["  Dark Mode Toggle", "Switches between light and dark themes"],
        ["  Language Selector", "Turkish, English, and other supported languages"],
        ["  'Delete Account' Button", "Account deletion with confirmation dialog"],
    ]

    # Build table
    table = doc.add_table(rows=1 + len(screen_list_rows), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(screen_list_headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "002B5C")

    # Data
    for r_idx, (item, func) in enumerate(screen_list_rows):
        cell0 = table.rows[r_idx + 1].cells[0]
        cell1 = table.rows[r_idx + 1].cells[1]

        is_screen_header = not item.startswith("  ")

        cell0.text = ""
        p0 = cell0.paragraphs[0]
        run0 = p0.add_run(item)
        run0.font.size = Pt(8)
        if is_screen_header:
            run0.bold = True
            run0.font.color.rgb = NAVY
            set_cell_shading(cell0, "E3F2FD")
            set_cell_shading(cell1, "E3F2FD")
        else:
            run0.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        cell1.text = ""
        p1 = cell1.paragraphs[0]
        run1 = p1.add_run(func)
        run1.font.size = Pt(8)
        if is_screen_header:
            run1.bold = True
            run1.font.color.rgb = NAVY

    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 5. SCREEN DESIGNS (WIREFRAMES)
    # ══════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Screen Designs (Wireframes)", 1)

    wf_text = (
        "Below are wireframe-level screen designs for TravAll's key screens. "
        "These wireframes follow the design principles established in Phase 2 and 3: "
        "mobile-first layout, thumb-zone ergonomics, card-based visual hierarchy, "
        "color-coded status indicators, and explicit user confirmation for critical actions. "
        "All wireframes are designed for a standard smartphone viewport (9:16 aspect ratio)."
    )
    add_body_text(doc, wf_text)

    wf_note = (
        "Note: The bottom navigation bar with 5 tabs (Home, Visa, Trips, Social, Profile) "
        "is persistent across all primary screens. The active tab is highlighted in blue."
    )
    p = doc.add_paragraph()
    run = p.add_run(wf_note)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # Generate wireframes
    wireframes = [
        ("5.1 Onboarding – Passport Selection", create_wireframe_onboarding,
         "The onboarding flow guides first-time users through passport setup. "
         "Large, clearly labeled passport options support Melih's use case (anxious first-timer). "
         "Progress dots indicate 3-step flow. 'Skip for now' respects user autonomy."),
        ("5.2 Dashboard", create_wireframe_dashboard,
         "The Dashboard is the main hub. A personalized greeting and upcoming trip card "
         "are at the top. Quick Action cards provide one-tap access to core features. "
         "The notification feed and Smart Suggestion card are below. The FAB allows quick trip creation."),
        ("5.3 Visa Assistant – Country List", create_wireframe_visa,
         "The Visa Assistant shows a passport-aware country list. Color-coded badges "
         "(green = visa-free, orange = eVisa, red = required) provide instant visual scanning. "
         "Filter chips and search bar allow efficient narrowing."),
        ("5.4 Visa Detail – Czech Republic", create_wireframe_visa_detail,
         "The Visa Detail screen shows full requirements, document checklist, and readiness progress. "
         "An orange alert banner highlights recent rule changes (Nihan's scenario). "
         "Upload and Guide buttons are prominently placed."),
        ("5.5 My Trips", create_wireframe_trips,
         "My Trips displays upcoming, past, and draft trips in color-coded cards. "
         "Each card summarizes flight and hotel info. The '+ Plan New Trip' button is prominent. "
         "A stats mini-card at bottom shows yearly travel summary."),
        ("5.6 Community Feed", create_wireframe_community,
         "The Community feed shows verified user posts with reactions. "
         "Authenticated badges (🔵 Verified, 🎓 Student) build trust. "
         "The FAB creates new posts. Tips tab provides curated travel advice."),
        ("5.7 Profile", create_wireframe_profile,
         "The Profile screen shows user identity, travel stats summary, and settings menu. "
         "Seven menu items are accessible with clear labels and icons. "
         "The layout follows one-handed ergonomics with tappable rows."),
    ]

    for i, (title, create_fn, desc) in enumerate(wireframes):
        if i > 0 and i % 1 == 0:
            doc.add_page_break()

        add_heading_styled(doc, title, 2)
        add_body_text(doc, desc, size=10)
        doc.add_paragraph()

        print(f"  Generating wireframe: {title}...")
        fig = create_fn()
        fig_to_docx_image(doc, fig, width=3.5)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Figure {i+2}: {title.split(' – ')[-1] if '–' in title else title.split(' ', 1)[-1]} Screen")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    # ══════════════════════════════════════════════════════════
    # SAVE
    # ══════════════════════════════════════════════════════════
    doc.save(DOC_PATH)
    print(f"\n✅ Phase 4 document saved to:\n   {DOC_PATH}")
    print(f"   File size: {os.path.getsize(DOC_PATH) / 1024:.0f} KB")


if __name__ == "__main__":
    print("=" * 60)
    print("  TravAll Phase 4: Generating DOCX Document")
    print("=" * 60)
    build_document()
